"""#########################################################################################################
 
 Script Name   :DocModder.py
 Description   :Takes in arguments, searches and replaces a .docx Word template. For onboarding new users.
 Author        :Andrew Campeau
 Last Compiled :2020-08-06
 Args          :-h, --help            show this help message and exit
                -e EMAIL, --email EMAIL
                -n NAME, --name NAME
                -u USERNAME, --username USERNAME
                -p PASSWORD, --password PASSWORD
                -t TEMPPATH, --tempPath TEMPPATH
                -s SAVEPATH, --savePath SAVEPATH

 Notes:
 Uses Pyinstaller to compile to .exe

#########################################################################################################"""

import docx
import argparse
import os
import sys
import time

def docx_find_replace_text(doc, search_text, replace_text):
    paragraphs = list(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs.append(paragraph)
    for p in paragraphs:
        if search_text in p.text:
            inline = p.runs
            # Replace strings and retain the same style.
            # The text to be replaced can be split over several runs so
            # search through, identify which runs need to have text replaced
            # then replace the text in those identified
            started = False
            search_index = 0
            # found_runs is a list of (inline index, index of match, length of match)
            found_runs = list()
            found_all = False
            replace_done = False
            for i in range(len(inline)):

                # case 1: found in single run so short circuit the replace
                if search_text in inline[i].text and not started:
                    found_runs.append((i, inline[i].text.find(search_text), len(search_text)))
                    text = inline[i].text.replace(search_text, str(replace_text))
                    inline[i].text = text
                    replace_done = True
                    found_all = True
                    break

                if search_text[search_index] not in inline[i].text and not started:
                    # keep looking ...
                    continue

                # case 2: search for partial text, find first run
                if search_text[search_index] in inline[i].text and inline[i].text[-1] in search_text and not started:
                    # check sequence
                    start_index = inline[i].text.find(search_text[search_index])
                    check_length = len(inline[i].text)
                    for text_index in range(start_index, check_length):
                        if inline[i].text[text_index] != search_text[search_index]:
                            # no match so must be false positive
                            break
                    if search_index == 0:
                        started = True
                    chars_found = check_length - start_index
                    search_index += chars_found
                    found_runs.append((i, start_index, chars_found))
                    if search_index != len(search_text):
                        continue
                    else:
                        # found all chars in search_text
                        found_all = True
                        break

                # case 2: search for partial text, find subsequent run
                if search_text[search_index] in inline[i].text and started and not found_all:
                    # check sequence
                    chars_found = 0
                    check_length = len(inline[i].text)
                    for text_index in range(0, check_length):
                        if inline[i].text[text_index] == search_text[search_index]:
                            search_index += 1
                            chars_found += 1
                        else:
                            break
                    # no match so must be end
                    found_runs.append((i, 0, chars_found))
                    if search_index == len(search_text):
                        found_all = True
                        break

            if found_all and not replace_done:
                for i, item in enumerate(found_runs):
                    index, start, length = [t for t in item]
                    if i == 0:
                        text = inline[index].text.replace(inline[index].text[start:start + length], str(replace_text))
                        inline[index].text = text
                    else:
                        text = inline[index].text.replace(inline[index].text[start:start + length], '')
                        inline[index].text = text
            # print(p.text)

my_parser = argparse.ArgumentParser()
my_parser.add_argument('-e', '--email', dest='email', action='store')
my_parser.add_argument('-n', '--name', dest='name', action='store')
my_parser.add_argument('-u', '--username', dest='username', action='store')
my_parser.add_argument('-p', '--password', dest='password', action='store')
my_parser.add_argument('-t', '--tempPath', dest='tempPath', action='store')
my_parser.add_argument('-s', '--savePath', dest='savePath', action='store')

results = my_parser.parse_args()
print(vars(results))

doc = docx.Document(results.tempPath)
docx_find_replace_text(doc, '[Email]', results.email)
docx_find_replace_text(doc, '[FullName]', results.name)
docx_find_replace_text(doc, '[Username]', results.username)
docx_find_replace_text(doc, '[Password]', results.password)
doc.save(results.savePath)